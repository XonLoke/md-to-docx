"""
Markdown to DOCX Converter
Converts markdown content into a nicely formatted Word document.

Supports:
  - Standard markdown: headings, bold/italic, code, links, lists, tables, etc.
  - Embedded HTML formatting: <b>, <i>, <code>, <img> alt text, <br>
  - HTML <table> blocks converted to DOCX tables
  - Nested inline formatting: **bold containing [links](url)** renders correctly
"""

import re
import io
from html.parser import HTMLParser
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import xml.etree.ElementTree as ET


# ── Colour palette ──────────────────────────────────────────────────────────
HEADING_COLOR   = RGBColor(0x1F, 0x49, 0x7D)   # dark blue
H1_COLOR        = RGBColor(0x1F, 0x49, 0x7D)
H2_COLOR        = RGBColor(0x2E, 0x74, 0xB5)
H3_COLOR        = RGBColor(0x2E, 0x74, 0xB5)
CODE_BG         = "F2F2F2"                       # light grey (shading)
CODE_FONT_COLOR = RGBColor(0xC7, 0x25, 0x4E)    # red-ish for inline code
BLOCKQUOTE_COLOR= RGBColor(0x59, 0x59, 0x59)
LINK_COLOR      = RGBColor(0x00, 0x56, 0xB3)
TABLE_HEADER_BG = "1F497D"                       # dark blue fill
TABLE_HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)    # white text
TABLE_ROW_ALT   = "DEEAF1"                       # light blue alternate row


#-------------------------------------------------------------------------------
# SECTION: XML / OXML Helpers
# Purpose: Low-level helper functions for manipulating DOCX XML elements
#          (shading, borders, cell backgrounds).
#-------------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str):
    """Apply a background fill colour to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_para_shading(para, hex_color: str):
    """Apply paragraph-level shading (used for code blocks)."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    pPr.append(shd)


def _set_para_border(para, side="left", color="2E74B5", space="4", sz="24"):
    """Add a coloured border to a paragraph (used for blockquotes)."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bd   = OxmlElement(f"w:{side}")
    bd.set(qn("w:val"),   "single")
    bd.set(qn("w:sz"),    sz)
    bd.set(qn("w:space"), space)
    bd.set(qn("w:color"), color)
    pBdr.append(bd)
    pPr.append(pBdr)


def _add_horizontal_rule(doc: Document):
    """Insert a full-width horizontal rule paragraph."""
    para = doc.add_paragraph()
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "AAAAAA")
    pBdr.append(bot)
    pPr.append(pBdr)
    return para


#-------------------------------------------------------------------------------
# SECTION: Text Cleaning - HTML to Plain Text / Markdown
# Purpose: Converts embedded HTML formatting tags to their markdown equivalents
#          so the inline parser can handle them. Strips or replaces remaining
#          HTML tags (images, comments, structural tags) with readable text.
#-------------------------------------------------------------------------------

def _clean_html_text(text: str) -> str:
    """Convert HTML formatting to markdown equivalents and strip remaining tags.

    This handles:
      - <b>/<strong> to **bold**  (markdown bold syntax)
      - <i>/<em>     to *italic*  (markdown italic syntax)
      - <code>       to inline code
      - <del>/<s>    to ~~strikethrough~~
      - <u>/<ins>    to underlined (best-effort)
      - <img>        to [Image: alt text]
      - <br>         to newline
      - All other HTML tags are stripped (content preserved)
      - HTML comments are removed entirely
    """
    # Remove HTML comments first (they can contain anything)
    text = re.sub(r'<!--.*?-->', '', text, flags=re.DOTALL)

    # Replace <br> and <br/> with newline (preserves line breaks)
    text = re.sub(r'<br\s*/?>', '\n', text)

    # Extract <img> alt text: prefer alt attribute, fall back to src
    def _replace_img(m):
        alt = m.group(1) or ""
        return f"[Image: {alt}]" if alt else "[Image]"
    text = re.sub(r'<img[^>]*alt=["\']([^"\']*)["\'][^>]*>', _replace_img, text)
    text = re.sub(r'<img[^>]*>', "[Image]", text)

    # Convert inline HTML formatting tags to their markdown equivalents.
    conversions = [
        ("<strong>", "**"), ("</strong>", "**"),
        ("<b>", "**"),      ("</b>", "**"),
        ("<em>", "*"),      ("</em>", "*"),
        ("<i>", "*"),       ("</i>", "*"),
        ("<code>", "`"),    ("</code>", "`"),
        ("<del>", "~~"),    ("</del>", "~~"),
        ("<s>", "~~"),      ("</s>", "~~"),
        ("<ins>", "__"),    ("</ins>", "__"),
        ("<u>", "__"),      ("</u>", "__"),
    ]
    for html_tag, md_syntax in conversions:
        text = text.replace(html_tag, md_syntax)

    # Strip all remaining HTML tags (e.g. <p>, <div>, <span>, <a>)
    # Keep their inner text content - only the tags are removed.
    text = re.sub(r'<[^>]+>', "", text)

    # Collapse multiple spaces into one (preserve single spaces)
    text = re.sub(r"  +", " ", text)

    return text.strip()


#-------------------------------------------------------------------------------
# SECTION: Run-level Formatting
# Purpose: Apply bold/italic/code/colour/strikethrough properties to a single
#          DOCX run. Called by the inline parser for each matched token.
#-------------------------------------------------------------------------------

def _apply_inline(run, bold=False, italic=False, code=False,
                  color: RGBColor = None, underline=False, strike=False):
    """Apply text formatting properties to a DOCX run."""
    run.bold      = bold
    run.italic    = italic
    run.underline = underline
    run.font.strike = strike
    if code:
        run.font.name        = "Courier New"
        run.font.size        = Pt(9.5)
        run.font.color.rgb   = CODE_FONT_COLOR
    if color:
        run.font.color.rgb = color


#-------------------------------------------------------------------------------
# SECTION: Inline Markdown Parser (Recursive)
# Purpose: Parse markdown inline formatting (**bold**, *italic*, inline code,
#          [links](url), ~~strikethrough~~, images) and add styled runs to
#          a paragraph. Recursively handles nested formatting, e.g.
#          **bold with [a link](url) inside** renders the link correctly.
#-------------------------------------------------------------------------------

_INLINE_RE = re.compile(
    r"(\*\*\*(?P<bolditalic>.+?)\*\*\*)"
    r"|(\*\*(?P<bold>.+?)\*\*)"
    r"|(__(?P<bold2>.+?)__)"
    r"|(\*(?P<italic>.+?)\*)"
    r"|(_(?P<italic2>.+?)_)"
    r"|(~~(?P<strike>.+?)~~)"
    r"|(`(?P<code>.+?)`)"
    r"|(\[(?P<link_text>[^\]]+)\]\((?P<link_url>[^\)]+)\))"
    r"|(\!\[(?P<img_alt>[^\]]*)\]\((?P<img_url>[^\)]+)\))",
    re.DOTALL
)


def _add_inline_text(para, text: str, base_bold=False, base_italic=False,
                     base_color: RGBColor = None, base_strike=False,
                     base_code=False):
    """Parse inline markdown in *text* and add styled runs to *para*.

    When bold, italic, or strikethrough patterns match, this function
    recurses into the matched content so nested formatting (e.g. links
    inside bold) is also parsed. Formatting properties passed as
    ``base_*`` are inherited by all runs created during this call.
    """
    pos = 0
    for m in _INLINE_RE.finditer(text):
        # Plain text before this match (with inherited base formatting)
        if m.start() > pos:
            run = para.add_run(text[pos:m.start()])
            _apply_inline(run, bold=base_bold, italic=base_italic,
                         color=base_color, code=base_code, strike=base_strike)

        g = m.groupdict()

        # ***bold italic***
        if g.get("bolditalic"):
            _add_inline_text(para, g["bolditalic"],
                           base_bold=True, base_italic=True,
                           base_color=base_color, base_strike=base_strike,
                           base_code=base_code)

        # **bold** or __bold__
        elif g.get("bold") or g.get("bold2"):
            content = g["bold"] or g["bold2"]
            _add_inline_text(para, content,
                           base_bold=True,
                           base_italic=base_italic,
                           base_color=base_color,
                           base_strike=base_strike,
                           base_code=base_code)

        # *italic* or _italic_
        elif g.get("italic") or g.get("italic2"):
            content = g["italic"] or g["italic2"]
            _add_inline_text(para, content,
                           base_bold=base_bold,
                           base_italic=True,
                           base_color=base_color,
                           base_strike=base_strike,
                           base_code=base_code)

        # ~~strikethrough~~
        elif g.get("strike"):
            _add_inline_text(para, g["strike"],
                           base_bold=base_bold,
                           base_italic=base_italic,
                           base_color=base_color,
                           base_strike=True,
                           base_code=base_code)

        # inline code
        elif g.get("code"):
            run = para.add_run(g["code"])
            _apply_inline(run, code=True,
                         bold=base_bold, italic=base_italic,
                         color=base_color, strike=base_strike)

        # [link text](url)
        elif g.get("link_text"):
            run = para.add_run(g["link_text"])
            _apply_inline(run, color=LINK_COLOR, underline=True,
                         bold=base_bold, italic=base_italic,
                         strike=base_strike)

        # ![alt](url) image
        elif g.get("img_alt") is not None:
            alt = g["img_alt"] or g["img_url"]
            run = para.add_run(f"[Image: {alt}]")
            run.italic = True
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        pos = m.end()

    # Remaining text after last match
    if pos < len(text):
        run = para.add_run(text[pos:])
        _apply_inline(run, bold=base_bold, italic=base_italic,
                     color=base_color, code=base_code, strike=base_strike)


#-------------------------------------------------------------------------------
# SECTION: Document Style Setup
# Purpose: Create a new blank Document with default margins and body font.
#-------------------------------------------------------------------------------

def _setup_document() -> Document:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(3.0)

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    return doc


#-------------------------------------------------------------------------------
# SECTION: Block-level Renderers
# Purpose: Render individual block-level elements (headings, paragraphs, code
#          blocks, blockquotes, lists, tables) into the DOCX document.
#          Each renderer cleans HTML from its text content before parsing
#          inline formatting.
#-------------------------------------------------------------------------------

def _render_heading(doc: Document, text: str, level: int):
    """Render a markdown heading (# to ######) as a styled DOCX heading."""
    text = _clean_html_text(text)
    style_name = f"Heading {min(level, 6)}"
    para = doc.add_paragraph(style=style_name)
    para.clear()

    colors = {1: H1_COLOR, 2: H2_COLOR, 3: H3_COLOR}
    color  = colors.get(level, HEADING_COLOR)
    sizes  = {1: 22, 2: 18, 3: 15, 4: 13, 5: 12, 6: 11}
    size   = sizes.get(level, 11)

    _add_inline_text(para, text, base_bold=(level <= 3), base_color=color)
    for run in para.runs:
        run.font.size = Pt(size)
        run.font.color.rgb = color
        if level == 1:
            run.bold = True
    return para


def _render_paragraph(doc: Document, text: str):
    """Render a plain markdown paragraph with inline formatting."""
    text = _clean_html_text(text)
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    _add_inline_text(para, text)
    return para


def _render_code_block(doc: Document, code: str, lang: str = ""):
    """Render a fenced code block with grey background and monospace font."""
    lines = code.split("\n")
    for i, line in enumerate(lines):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(2) if i > 0 else Pt(6)
        para.paragraph_format.space_after  = Pt(2) if i < len(lines) - 1 else Pt(6)
        para.paragraph_format.left_indent  = Cm(0.5)
        _set_para_shading(para, CODE_BG)
        run = para.add_run(line if line else " ")
        run.font.name  = "Courier New"
        run.font.size  = Pt(9.5)
        run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)


def _render_blockquote(doc: Document, text: str):
    """Render a blockquote with left blue border accent and italic text."""
    text = _clean_html_text(text)
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Cm(1.0)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    _set_para_border(para)
    _add_inline_text(para, text, base_italic=True, base_color=BLOCKQUOTE_COLOR)
    return para


def _render_list_item(doc: Document, text: str, level: int, ordered: bool,
                      counter: int = 1):
    """Render a single list item (bullet or numbered)."""
    text = _clean_html_text(text)
    style = "List Number" if ordered else "List Bullet"
    para  = doc.add_paragraph(style=style)
    para.paragraph_format.left_indent   = Cm(0.5 * (level + 1))
    para.paragraph_format.space_after   = Pt(3)
    _add_inline_text(para, text)
    return para


def _render_table(doc: Document, rows: list):
    """Render a table with styled header row and alternating row colours.

    Both header and body rows use _add_inline_text for inline formatting
    (bold/italic/code/links inside cells).
    """
    if not rows:
        return

    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx in range(col_count):
            cell = row.cells[c_idx]
            text = row_data[c_idx] if c_idx < len(row_data) else ""
            text = text.strip()

            if r_idx == 0:
                # Header row: dark blue background, white, bold
                _set_cell_bg(cell, TABLE_HEADER_BG)
                para = cell.paragraphs[0]
                _add_inline_text(para, text)
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = TABLE_HEADER_FG
                    run.font.size = Pt(10)
            else:
                # Body row: alternating shading, inline formatting
                if r_idx % 2 == 0:
                    _set_cell_bg(cell, TABLE_ROW_ALT)
                para = cell.paragraphs[0]
                _add_inline_text(para, text)
                for run in para.runs:
                    run.font.size = Pt(10)

    # Auto-fit columns
    for col in table.columns:
        for cell in col.cells:
            cell.width = Inches(6.0 / col_count)

    doc.add_paragraph()  # spacing after table


#-------------------------------------------------------------------------------
# SECTION: HTML Table Parser
# Purpose: Parse an HTML <table> block and return a list of rows (each row
#          is a list of cell text strings) suitable for _render_table().
#-------------------------------------------------------------------------------

class _HTMLTableParser(HTMLParser):
    """Minimal HTML <table> parser. Extracts row/cell content as plain text.

    Handles:
      - <table>, <tr>, <td>, <th> tags
      - Nested inline tags (<b>, <i>, <a>) - inner text is collected
      - <br> tags within cells to newlines
    Ignores: colspan, rowspan, and all other HTML attributes.
    """

    def __init__(self):
        super().__init__()
        self.rows: list[list[str]] = []
        self._current_row: list[str] = []
        self._current_cell: list[str] = []
        self._in_cell = False

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]):
        tag = tag.lower()
        if tag in ("td", "th"):
            self._in_cell = True
            self._current_cell = []
        elif tag == "tr":
            self._current_row = []
        elif tag == "br":
            if self._in_cell:
                self._current_cell.append("\n")

    def handle_endtag(self, tag: str):
        tag = tag.lower()
        if tag in ("td", "th"):
            self._in_cell = False
            cell_text = "".join(self._current_cell).strip()
            self._current_row.append(cell_text)
            self._current_cell = []
        elif tag == "tr":
            if self._current_row:
                self.rows.append(self._current_row)
            self._current_row = []
        elif tag == "table":
            if self._current_row:
                self.rows.append(self._current_row)
                self._current_row = []

    def handle_data(self, data: str):
        if self._in_cell:
            self._current_cell.append(data)


def _render_html_table(doc: Document, html: str):
    """Parse an HTML <table> block and render it as a DOCX table."""
    parser = _HTMLTableParser()
    try:
        parser.feed(html)
    except Exception:
        return
    if parser.rows:
        _render_table(doc, parser.rows)


#-------------------------------------------------------------------------------
# SECTION: Main Parser / State Machine
# Purpose: Core markdown-to-DOCX conversion engine. Walks lines sequentially,
#          detecting block-level elements (headings, code fences, lists,
#          tables, blockquotes, horizontal rules, HTML tables) and dispatching
#          to the appropriate renderer.
#-------------------------------------------------------------------------------

def convert_markdown_to_docx(markdown_text: str) -> bytes:
    """
    Convert a markdown string to a DOCX file.

    Args:
        markdown_text: Raw markdown content (may include HTML).

    Returns:
        Raw bytes of the generated .docx file.
    """
    doc = _setup_document()
    lines = markdown_text.replace("\r\n", "\n").split("\n")

    i = 0
    in_code_block  = False
    code_lines     = []
    code_lang      = ""
    in_table       = False
    table_rows     = []

    ol_counters: dict[int, int] = {}

    def flush_table():
        nonlocal in_table, table_rows
        if table_rows:
            clean = [r for r in table_rows
                     if not all(re.match(r"^:?-+:?$", c.strip()) for c in r if c.strip())]
            _render_table(doc, clean)
        in_table   = False
        table_rows = []

    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.strip().startswith("```"):
            if in_table:
                flush_table()
            if not in_code_block:
                in_code_block = True
                code_lang     = line.strip()[3:].strip()
                code_lines    = []
            else:
                in_code_block = False
                _render_code_block(doc, "\n".join(code_lines), code_lang)
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^\s*([-*_])\s*(\1\s*){2,}$", line):
            if in_table:
                flush_table()
            _add_horizontal_rule(doc)
            i += 1
            continue

        # Headings (ATX style)
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            if in_table:
                flush_table()
            level = len(heading_match.group(1))
            text  = heading_match.group(2).strip()
            _render_heading(doc, text, level)
            i += 1
            continue

        # Setext headings
        if i + 1 < len(lines):
            next_line = lines[i + 1]
            if re.match(r"^=+\s*$", next_line) and line.strip():
                if in_table:
                    flush_table()
                _render_heading(doc, line.strip(), 1)
                i += 2
                continue
            if re.match(r"^-+\s*$", next_line) and line.strip():
                if in_table:
                    flush_table()
                _render_heading(doc, line.strip(), 2)
                i += 2
                continue

        # Blockquote
        bq_match = re.match(r"^>\s?(.*)", line)
        if bq_match:
            if in_table:
                flush_table()
            bq_lines = [bq_match.group(1)]
            while i + 1 < len(lines) and re.match(r"^>\s?(.*)", lines[i + 1]):
                i += 1
                bq_lines.append(re.match(r"^>\s?(.*)", lines[i]).group(1))
            _render_blockquote(doc, " ".join(bq_lines))
            i += 1
            continue

        # Unordered list
        ul_match = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
        if ul_match:
            if in_table:
                flush_table()
            indent = len(ul_match.group(1)) // 2
            _render_list_item(doc, ul_match.group(2), indent, ordered=False)
            i += 1
            continue

        # Ordered list
        ol_match = re.match(r"^(\s*)(\d+)\.\s+(.+)$", line)
        if ol_match:
            if in_table:
                flush_table()
            indent  = len(ol_match.group(1)) // 2
            counter = ol_counters.get(indent, 0) + 1
            ol_counters[indent] = counter
            _render_list_item(doc, ol_match.group(3), indent, ordered=True,
                              counter=counter)
            i += 1
            continue
        else:
            if not re.match(r"^\s*$", line):
                ol_counters = {}

        # HTML <table> block
        if re.match(r"^\s*<table", line, re.IGNORECASE):
            if in_table:
                flush_table()
            html_lines = [line]
            i += 1
            while i < len(lines) and not re.match(r"^\s*</table", lines[i], re.IGNORECASE):
                html_lines.append(lines[i])
                i += 1
            if i < len(lines):
                html_lines.append(lines[i])
            _render_html_table(doc, "\n".join(html_lines))
            i += 1
            continue

        # Pipe table
        if "|" in line:
            cells = [c for c in line.split("|")]
            if cells and cells[0].strip() == "":
                cells = cells[1:]
            if cells and cells[-1].strip() == "":
                cells = cells[:-1]
            if cells:
                in_table = True
                table_rows.append(cells)
                i += 1
                continue

        if in_table:
            flush_table()

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        para_lines = [line]
        while (i + 1 < len(lines)
               and lines[i + 1].strip()
               and not re.match(r"^(#{1,6}\s|[-*+]\s|\d+\.\s|>|\s*```|\|)", lines[i + 1])):
            i += 1
            para_lines.append(lines[i])
        _render_paragraph(doc, " ".join(para_lines))
        i += 1

    if in_table:
        flush_table()
    if in_code_block and code_lines:
        _render_code_block(doc, "\n".join(code_lines), code_lang)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
